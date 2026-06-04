"""
PDF to Excel BOM Extraction Tool

This module handles extraction of Bill of Materials (BOM) data from digitally generated PDFs
and converts them to clean Excel files compatible with the BOM Comparison tool.

Key Features:
- Validates digital PDFs (rejects scanned/image-only PDFs)
- Detects BOM tables with confidence scoring
- Normalizes extracted data (handles merged cells, multi-line items, duplicate headers)
- Maps to standardized BOM schema (MPN, Qty, Ref Des, Description)
- Generates Excel with no merged cells
"""

import pdfplumber
import pypdf
import pandas as pd
import tempfile
import io
import os
from typing import Dict, List, Tuple, Optional
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from excel_tool import find_header_row_and_map
from ai_helpers import (
    classify_column,
    score_column_content,
    is_mpn_like,
)


# Error messages for graceful failures
ERROR_MESSAGES = {
    'scanned_pdf': (
        "This PDF appears to be a scanned document or image-only. "
        "Please provide a digitally generated PDF with selectable text."
    ),
    'no_table': (
        "Could not detect a BOM table in this PDF. "
        "Ensure the PDF contains a table with Part Number, Quantity, and Description columns."
    ),
    'no_mpn': (
        "No part number column found. "
        "BOM must include a column for MPN, Part Number, P/N, or similar."
    ),
    'corrupt': (
        "PDF file is corrupted or invalid. Please try re-exporting from source application."
    ),
    'encrypted': (
        "Password-protected PDFs are not supported. Please provide unencrypted PDF."
    ),
    'low_confidence': (
        "BOM table detected with low confidence. Results may require manual review."
    )
}


# BOM header keywords with weights for table detection
# Enhanced to support: Manufacturer, MPN, Qty, Description, RefDes, Vendor, Vendor P/N, Item #, Customer P/N
HEADER_KEYWORDS = {
    # MPN / Part Number keywords
    'mpn': 10, 'part number': 10, 'mfg pn': 10, 'p/n': 8, 'part #': 10,
    'manufacturer part number': 10, 'mfg part number': 10, 'model number': 10,
    'part no': 9, 'part_number': 9, 'mfr pn': 10, 'mfr part': 9,

    # Manufacturer keywords
    'manufacturer': 9, 'mfg': 8, 'mfr': 8, 'vendor': 7, 'supplier': 7,
    'manufacturer name': 9, 'mfg name': 8,

    # Vendor/Supplier keywords
    'vendor part number': 8, 'supplier part number': 8, 'vendor pn': 7,
    'supplier pn': 7, 'vendor p/n': 7, 'supplier p/n': 7,

    # Customer Part Number keywords
    'customer pn': 8, 'customer part number': 8, 'customer p/n': 7,
    'internal pn': 7, 'company pn': 7, 'kjr p/n': 9,

    # Quantity keywords
    'quantity': 8, 'qty': 8, 'count': 5, 'amount': 5, 'qty required': 7,
    'qty per': 7, 'qty/assy': 7, 'per assy': 6,

    # Reference Designator keywords
    'refdes': 7, 'ref des': 7, 'designator': 7, 'location': 6, 'loc': 6,
    'reference designator': 8, 'reference designation': 7, 'refdes/loc': 7,
    'ref desg': 7, 'reference': 6,

    # Description keywords
    'description': 8, 'desc': 7, 'notes': 5, 'comment': 5, 'remarks': 4,
    'part description': 8, 'component description': 8, 'comments': 5,

    # Item/Line keywords
    'item': 6, 'line': 4, 'item number': 7, 'item #': 6, 'line number': 5,
    'line #': 4, 'rev': 3, 'revision': 3
}


def validate_digital_pdf(pdf_path: str) -> None:
    """
    Validate PDF is digitally generated (not scanned).

    Args:
        pdf_path: Path to PDF file

    Raises:
        ValueError: If PDF is scanned/image-only, encrypted, corrupt, or has no pages
    """
    try:
        # Check for encryption using pypdf
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                if pdf_reader.is_encrypted:
                    raise ValueError(ERROR_MESSAGES['encrypted'])
        except Exception as e:
            if "encrypted" in str(e).lower():
                raise ValueError(ERROR_MESSAGES['encrypted'])

        # Check for text content using pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            if len(pdf.pages) == 0:
                raise ValueError("PDF has no pages")

            # Check first 3 pages for text content
            text_found = False
            for page in pdf.pages[:3]:
                text = page.extract_text()
                # Lower threshold to 20 chars to catch minimal text PDFs
                if text and len(text.strip()) > 20:
                    text_found = True
                    break

            if not text_found:
                raise ValueError(ERROR_MESSAGES['scanned_pdf'])

    except ValueError:
        # Re-raise our custom validation errors
        raise
    except Exception as e:
        if "encrypted" in str(e).lower():
            raise ValueError(ERROR_MESSAGES['encrypted'])
        # For other errors, assume corrupt PDF
        raise ValueError(ERROR_MESSAGES['corrupt'])


def _headers_match(header1: List, header2: List) -> bool:
    """
    Check if two headers represent the same table structure.

    Args:
        header1: First header row
        header2: Second header row

    Returns:
        True if headers match with >= 80% similarity
    """
    if len(header1) != len(header2):
        return False

    matches = sum(1 for h1, h2 in zip(header1, header2)
                  if str(h1).lower().strip() == str(h2).lower().strip())

    return matches / len(header1) >= 0.8 if header1 else False


def _extract_bom_by_word_positions(pdf_path: str) -> Optional[Dict]:
    """
    Extract BOM data using word-level positions from pdfplumber.

    This is the most accurate method as it uses the actual x-coordinates
    of words in the PDF to determine column boundaries.
    """
    import re

    bom_keywords = [
        'item', 'number', 'revision', 'rev', 'description', 'desc',
        'lifecycle', 'phase', 'quantity', 'qty', 'manufacturer', 'mfg', 'mfr',
        'part', 'mpn', 'p/n', 'distributor', 'vendor', 'supplier',
        'reference', 'refdes', 'designator', 'comments', 'notes', 'model',
        'bom', 'cost', 'assy', 'per'
    ]

    with pdfplumber.open(pdf_path) as pdf:
        all_rows = []
        header_cols = None
        header_y = None

        for page_num, page in enumerate(pdf.pages):
            words = page.extract_words()
            if not words:
                continue

            # Group words by y-position (same line)
            lines_dict = {}
            for word in words:
                # Round y to group words on same line (tolerance of 8 units for varying baselines)
                y_key = round(word['top'] / 8) * 8
                if y_key not in lines_dict:
                    lines_dict[y_key] = []
                lines_dict[y_key].append(word)

            # Sort lines by y position
            sorted_y = sorted(lines_dict.keys())

            for y in sorted_y:
                line_words = sorted(lines_dict[y], key=lambda w: w['x0'])

                # Check if this is a header line
                line_text = ' '.join(w['text'].lower() for w in line_words)
                keyword_count = sum(1 for kw in bom_keywords if kw in line_text)

                if keyword_count >= 3 and header_cols is None:
                    # This is likely the header - extract column boundaries
                    # First, calculate the median gap between words to determine threshold
                    gaps = []
                    for i in range(1, len(line_words)):
                        gap = line_words[i]['x0'] - line_words[i-1]['x1']
                        gaps.append(gap)

                    if gaps:
                        # Use median gap * 1.5 as threshold for column break
                        sorted_gaps = sorted(gaps)
                        median_gap = sorted_gaps[len(sorted_gaps) // 2]
                        gap_threshold = max(median_gap * 1.5, 15)  # At least 15 units
                    else:
                        gap_threshold = 20

                    header_cols = []
                    current_col = []
                    last_x1 = -100

                    for word in line_words:
                        # If gap exceeds threshold, start new column
                        if word['x0'] - last_x1 > gap_threshold and current_col:
                            col_name = ' '.join(w['text'] for w in current_col)
                            col_x0 = current_col[0]['x0']
                            header_cols.append({'name': col_name, 'x0': col_x0})
                            current_col = []

                        current_col.append(word)
                        last_x1 = word['x1']

                    # Don't forget last column
                    if current_col:
                        col_name = ' '.join(w['text'] for w in current_col)
                        col_x0 = current_col[0]['x0']
                        header_cols.append({'name': col_name, 'x0': col_x0})

                    # Add end boundaries - use midpoint between columns
                    for i, col in enumerate(header_cols):
                        if i + 1 < len(header_cols):
                            # Use midpoint between this column's last word and next column's first word
                            col['x1'] = (col['x0'] + header_cols[i + 1]['x0']) / 2 + (header_cols[i + 1]['x0'] - col['x0']) / 2
                            col['x1'] = header_cols[i + 1]['x0'] - 5  # Small buffer before next column
                        else:
                            col['x1'] = 10000  # Last column extends to end

                    header_y = y
                    continue

                # If we have headers, extract data row
                if header_cols and line_words:
                    # Skip if this looks like another header
                    if keyword_count >= 3:
                        continue

                    # Skip header continuation lines (close to header y)
                    if header_y and abs(y - header_y) < 15:
                        continue

                    row_data = {col['name']: '' for col in header_cols}

                    for word in line_words:
                        # Find which column this word belongs to based on x position
                        word_x = word['x0']
                        best_col = None
                        best_dist = float('inf')

                        for col in header_cols:
                            # Calculate distance to column center
                            col_center = (col['x0'] + col['x1']) / 2
                            dist = abs(word_x - col['x0'])

                            # Word belongs to column if it starts within column boundaries
                            if col['x0'] - 10 <= word_x < col['x1'] + 10:
                                if dist < best_dist:
                                    best_dist = dist
                                    best_col = col

                        if best_col:
                            if row_data[best_col['name']]:
                                row_data[best_col['name']] += ' ' + word['text']
                            else:
                                row_data[best_col['name']] = word['text']

                    # Only add if row has meaningful data
                    non_empty = sum(1 for v in row_data.values() if v.strip())
                    if non_empty >= 2:
                        all_rows.append(row_data)

        if not header_cols or not all_rows:
            return None

        # Create DataFrame
        col_names = [col['name'] for col in header_cols]
        df = pd.DataFrame(all_rows, columns=col_names)

        # Filter rows that look like data (not footer/notes)
        def is_data_row(row):
            values = [str(v).strip() for v in row if str(v).strip()]
            if len(values) < 2:
                return False
            # Check for part-number-like values (alphanumeric)
            for v in values:
                if len(v) >= 3 and (any(c.isalpha() for c in v) or any(c.isdigit() for c in v)):
                    return True
            return False

        df = df[df.apply(is_data_row, axis=1)]

        if df.empty:
            return None

        return {
            'table_data': df,
            'confidence': 0.75,  # Higher confidence for word-position method
            'page_numbers': list(range(1, len(pdf.pages) + 1)),
            'header_row': col_names
        }


def _detect_bom_columns_in_header(header_line: str) -> List[Tuple[int, int, str]]:
    """
    Detect BOM column names in a header line and return their positions.

    Returns list of (start_pos, end_pos, column_name) tuples.
    """
    import re

    # Known BOM column names in order of specificity (longest first to avoid partial matches)
    known_columns = [
        'manufacturer part number', 'manufacturer part no', 'manufacturer pn',
        'mfg part number', 'mfr part number', 'vendor part number', 'supplier part number',
        'reference designator', 'reference designation', 'ref designator',
        'lifecycle phase', 'part number', 'part no', 'part #', 'item number', 'item no',
        'item #', 'line number', 'model number', 'model no', 'customer pn',
        'manufacturer', 'description', 'quantity', 'revision', 'comments',
        'distributor', 'supplier', 'vendor', 'refdes', 'ref des', 'phase',
        'item', 'mpn', 'mfg', 'mfr', 'qty', 'rev', 'bom', 'assy', 'cost', 'type',
        'part', 'number', 'per assy', 'uom', 'designator', 'location', 'notes'
    ]

    header_lower = header_line.lower()
    found_columns = []

    # Track which positions have been claimed
    claimed = [False] * len(header_line)

    for col_name in known_columns:
        # Find all occurrences
        start = 0
        while True:
            pos = header_lower.find(col_name, start)
            if pos == -1:
                break

            # Check if this position overlaps with already claimed positions
            end_pos = pos + len(col_name)
            overlap = any(claimed[i] for i in range(pos, min(end_pos, len(claimed))))

            if not overlap:
                # Mark as claimed
                for i in range(pos, min(end_pos, len(claimed))):
                    claimed[i] = True

                # Get the actual text (preserve case)
                actual_name = header_line[pos:end_pos].strip()
                found_columns.append((pos, end_pos, actual_name))

            start = pos + 1

    # Sort by position
    found_columns.sort(key=lambda x: x[0])

    # Adjust end positions to be the start of next column
    result = []
    for i, (start, end, name) in enumerate(found_columns):
        if i + 1 < len(found_columns):
            actual_end = found_columns[i + 1][0]
        else:
            actual_end = len(header_line)
        result.append((start, actual_end, name))

    return result


def _extract_by_column_positions(line: str, col_positions: List[Tuple[int, int, str]]) -> List[str]:
    """
    Extract data from line based on column positions.
    """
    values = []
    padded_line = line.ljust(max(end for _, end, _ in col_positions) if col_positions else len(line))

    for start, end, _ in col_positions:
        value = padded_line[start:end].strip()
        values.append(value)

    return values


def _extract_table_from_text(pdf_path: str) -> Optional[Dict]:
    """
    Fallback: Extract BOM data from text when table extraction fails.

    This handles cases where PDF tables don't have visible gridlines
    and pdfplumber can't detect column boundaries.

    Uses multiple strategies:
    1. Position-based extraction using header column positions
    2. Multi-space splitting for simple cases
    3. Pattern matching for common BOM data types

    Returns:
        Dictionary with table_data, confidence, page_numbers, header_row
        or None if text extraction also fails
    """
    import re

    with pdfplumber.open(pdf_path) as pdf:
        all_text = ""
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                all_text += text + "\n"

        if not all_text.strip():
            return None

        lines = all_text.strip().split('\n')

        # Extended BOM keywords - ordered by priority/specificity
        bom_keywords = [
            'manufacturer part number', 'mfg part number', 'mfr part number',
            'part number', 'part no', 'part #', 'part#', 'p/n',
            'manufacturer', 'mfg', 'mfr',
            'description', 'desc',
            'qty', 'quantity', 'per assy',
            'ref des', 'refdes', 'reference designator', 'designator',
            'model number', 'model no', 'model',
            'revision', 'rev',
            'item', 'bom'
        ]

        # Try to find header line with BOM keywords
        header_line_idx = None
        header_keywords_found = []
        best_header_score = 0

        for idx, line in enumerate(lines[:40]):  # Check first 40 lines
            line_lower = line.lower()
            keywords_in_line = []

            for kw in bom_keywords:
                if kw in line_lower:
                    # Don't double count overlapping keywords
                    already_found = any(kw in existing or existing in kw for existing in keywords_in_line)
                    if not already_found:
                        keywords_in_line.append(kw)

            # Score this line (prefer more keywords, and longer/more specific keywords)
            score = sum(len(kw) for kw in keywords_in_line)

            # At least 2 BOM keywords with good score
            if len(keywords_in_line) >= 2 and score > best_header_score:
                header_line_idx = idx
                header_keywords_found = keywords_in_line
                best_header_score = score

        if header_line_idx is None:
            return None

        header_line = lines[header_line_idx]

        # Strategy 1: Detect known BOM column names in header
        col_positions = _detect_bom_columns_in_header(header_line)

        # Strategy 2: Multi-space splitting as fallback
        header_parts = re.split(r'\s{2,}|\t', header_line)
        header_parts = [h.strip() for h in header_parts if h.strip()]

        # Choose strategy based on quality
        use_position_based = len(col_positions) >= 3 and len(col_positions) <= 15

        if use_position_based:
            # Use detected column names
            header_parts = [name for _, _, name in col_positions]

        if len(header_parts) < 2:
            # Try single-space splitting if multi-space fails
            header_parts = header_line.split()
            if len(header_parts) < 3:
                return None
            use_position_based = False

        data_rows = []

        for line in lines[header_line_idx + 1:]:
            if not line.strip():
                continue

            # Skip lines that look like headers repeated
            line_lower = line.lower()
            header_match_count = sum(1 for kw in header_keywords_found if kw in line_lower)
            if header_match_count >= min(3, len(header_keywords_found)):
                continue

            # Skip lines that are too short
            if len(line.strip()) < 8:
                continue

            # Skip footer/header lines
            skip_patterns = ['page ', 'sheet ', 'revision date', 'date:', 'drawn by',
                           'approved by', 'copyright', 'confidential', 'total:']
            if any(skip in line_lower for skip in skip_patterns):
                continue

            # Skip lines that look like continuation markers or notes
            if line.strip().startswith('---') or line.strip().startswith('==='):
                continue

            # Skip underline rows (common in text-based tables)
            if re.match(r'^[\s_\-=]+$', line):
                continue

            # Extract data
            if use_position_based:
                parts = _extract_by_column_positions(line, col_positions)
            else:
                # Multi-space splitting
                parts = re.split(r'\s{2,}|\t', line)
                parts = [p.strip() for p in parts if p.strip()]

                # If splitting gives too few parts, try single space with smart merging
                if len(parts) < 2:
                    parts = line.split()

                    # Try to merge parts based on expected data patterns
                    # (e.g., descriptions are usually at the end and can have spaces)
                    if len(parts) > len(header_parts):
                        # Assume extra parts belong to description (usually last)
                        merged = parts[:len(header_parts)-1]
                        merged.append(' '.join(parts[len(header_parts)-1:]))
                        parts = merged

            # Remove empty parts from ends but preserve structure
            while parts and not parts[-1]:
                parts.pop()
            while parts and not parts[0]:
                parts.pop(0)

            # Validate: at least 2 non-empty parts
            non_empty = sum(1 for p in parts if p)
            if non_empty < 2:
                continue

            # Pad or trim to match header length
            target_len = len(header_parts)
            while len(parts) < target_len:
                parts.append('')
            if len(parts) > target_len:
                # Merge extra into last column (usually description)
                extra = ' '.join(parts[target_len:])
                parts = parts[:target_len]
                if parts[-1]:
                    parts[-1] = parts[-1] + ' ' + extra
                else:
                    parts[-1] = extra

            data_rows.append(parts)

        if not data_rows:
            return None

        # Create DataFrame
        df = pd.DataFrame(data_rows, columns=header_parts)

        # Filter rows that have at least one value that looks like a part number
        # (alphanumeric with at least some letters and numbers)
        def has_pn_like_value(row):
            for val in row:
                val_str = str(val).strip()
                if len(val_str) >= 3:
                    has_letters = any(c.isalpha() for c in val_str)
                    has_numbers = any(c.isdigit() for c in val_str)
                    if has_letters or has_numbers:
                        return True
            return False

        df = df[df.apply(has_pn_like_value, axis=1)]

        if df.empty:
            return None

        return {
            'table_data': df,
            'confidence': 0.6,  # Lower confidence for text extraction
            'page_numbers': list(range(1, len(pdf.pages) + 1)),
            'header_row': header_parts
        }


def _is_table_data_valid(table: List) -> bool:
    """
    Check if extracted table has valid data distribution.

    Returns False if all data is crammed into one column (extraction failed).
    """
    if not table or len(table) < 2:
        return False

    header = table[0]
    data_rows = table[1:]

    # Filter out completely empty rows
    non_empty_rows = []
    for row in data_rows:
        if any(cell and str(cell).strip() for cell in row):
            non_empty_rows.append(row)

    if not non_empty_rows:
        return False

    # Count how many columns actually have data across non-empty rows
    cols_with_data = 0
    total_cells_with_data = 0

    for col_idx in range(len(header)):
        col_has_data = False
        col_data_count = 0
        for row in non_empty_rows[:10]:  # Check first 10 non-empty rows
            if col_idx < len(row) and row[col_idx] and str(row[col_idx]).strip():
                col_has_data = True
                col_data_count += 1
        if col_has_data:
            cols_with_data += 1
            total_cells_with_data += col_data_count

    # If only 1-2 columns have data but header suggests many more, extraction failed
    if cols_with_data <= 2 and len(header) > 5:
        return False

    # Check if first column has all the data concatenated (Zoetis pattern)
    # This happens when pdfplumber can't detect column boundaries
    if cols_with_data == 1 and len(header) > 3:
        # Check if the single column with data has very long values
        # (which would indicate concatenated data)
        first_col_idx = 0
        for col_idx in range(len(header)):
            for row in non_empty_rows[:5]:
                if col_idx < len(row) and row[col_idx] and str(row[col_idx]).strip():
                    first_col_idx = col_idx
                    break
            if first_col_idx > 0:
                break

        avg_len = 0
        count = 0
        for row in non_empty_rows[:5]:
            if first_col_idx < len(row) and row[first_col_idx]:
                avg_len += len(str(row[first_col_idx]))
                count += 1

        if count > 0:
            avg_len = avg_len / count
            # If average cell length > 50 chars, it's likely concatenated data
            if avg_len > 50:
                return False

    return True


def _infer_page_rotation(page) -> int:
    """
    Best-effort guess of page rotation when /Rotate is missing.

    Many engineering drawings store landscape content on a portrait page
    (8.5x11) and rely on the renderer to know which way is up. We sniff the
    word orientations: if the dominant `upright` flag in extract_words is
    False, the content is sideways and we should rotate 90°. Cheap and
    conservative — only returns 90/270 if there's strong signal.
    """
    try:
        words = page.extract_words(extra_attrs=["upright"]) or []
    except Exception:
        return 0
    if not words:
        return 0
    sideways = sum(1 for w in words if w.get("upright") is False)
    if sideways > len(words) * 0.6:
        return 90
    return 0


def detect_bom_table(pdf_path: str) -> Dict:
    """
    Detect the primary BOM table using multi-factor scoring.

    Args:
        pdf_path: Path to PDF file

    Returns:
        Dictionary with:
            - table_data: pd.DataFrame with extracted data
            - confidence: float (0-1) confidence score
            - page_numbers: List of page numbers where table was found
            - header_row: List of header column names

    Raises:
        ValueError: If no BOM table detected with sufficient confidence
    """
    candidates = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Auto-rotate landscape engineering-drawing pages so word coords
            # come out in reading order. pdfplumber inherits the PDF's
            # /Rotate value but some drawings store rotation as 0 even
            # though the content is sideways; in that case, detecting
            # rotation here costs little (cheap heuristic — see helper).
            rotation = getattr(page, "rotation", 0) or _infer_page_rotation(page)
            if rotation in (90, 180, 270):
                try:
                    page = page.rotate(360 - rotation)  # bring back to upright
                except Exception:
                    pass

            tables = page.extract_tables()

            for table_idx, table in enumerate(tables):
                if not table or len(table) < 2:  # Need header + at least 1 data row
                    continue

                # Check if table data is valid (not all crammed into one column)
                if not _is_table_data_valid(table):
                    continue

                # Score this table
                score = 0
                header_row = table[0]

                # Factor 1: Header keyword matching
                header_text = ' '.join([str(cell).lower() if cell else '' for cell in header_row])
                for keyword, weight in HEADER_KEYWORDS.items():
                    if keyword in header_text:
                        score += weight

                # Factor 2: Row count (more rows = likely BOM)
                row_count = len(table)
                score += min(row_count / 5, 10)  # Cap at 10 points

                # Factor 3: Column count (typical BOM has 4-8 columns)
                col_count = len(header_row)
                if 4 <= col_count <= 8:
                    score += 5

                # Factor 4: Numeric data presence (Qty column indicator)
                numeric_cells = sum(1 for row in table[1:] for cell in row
                                  if cell and str(cell).replace('.','').replace(',','').isdigit())
                score += min(numeric_cells / len(table), 5)

                candidates.append({
                    'table': table,
                    'score': score,
                    'page': page_num + 1,
                    'table_idx': table_idx,
                    'header': header_row
                })

    # Sort by score descending
    candidates.sort(key=lambda x: x['score'], reverse=True)

    def _has_valid_mpn_column(df):
        """Check if DataFrame has a column that can be mapped to MPN."""
        for col in df.columns:
            col_type = detect_column_type(str(col))
            if col_type == 'mpn':
                return True
        return False

    if not candidates or candidates[0]['score'] < 15:  # Minimum confidence threshold
        # Try word-position extraction first (most accurate)
        word_result = _extract_bom_by_word_positions(pdf_path)
        if word_result and len(word_result['table_data']) > 0:
            # Verify it has valid columns before accepting
            if _has_valid_mpn_column(word_result['table_data']):
                return word_result

        # Then try text-based extraction as fallback
        text_result = _extract_table_from_text(pdf_path)
        if text_result and _has_valid_mpn_column(text_result['table_data']):
            return text_result

        # Return word result anyway if it exists (let downstream handle errors)
        if word_result and len(word_result['table_data']) > 0:
            return word_result
        if text_result:
            return text_result

        raise ValueError(ERROR_MESSAGES['no_table'])

    # Even if we found tables, verify the best one has valid data distribution
    # If not, try word-position or text extraction (handles Zoetis-style PDFs
    # where table structure is detected but data is all concatenated into one column)
    if not _is_table_data_valid(candidates[0]['table']):
        # Try word-position extraction first
        word_result = _extract_bom_by_word_positions(pdf_path)
        if word_result and len(word_result['table_data']) > 0:
            if _has_valid_mpn_column(word_result['table_data']):
                return word_result

        # Then try text extraction
        text_result = _extract_table_from_text(pdf_path)
        if text_result and _has_valid_mpn_column(text_result['table_data']):
            return text_result

        # Fall through to use the table anyway if all extraction methods fail

    # Take highest-scoring table
    best_table = candidates[0]
    best_col_count = len(best_table['header']) if best_table['header'] else 0

    # Multi-page continuation merge — be generous so we capture all rows,
    # even when continuation pages don't repeat the header.
    # Strategy:
    #   1) Same columns + matching header → strip the duplicate header, append data.
    #   2) Same columns + NO matching header → treat the entire continuation
    #      as data (no row to strip) — this catches Page 2..N of long BOMs that
    #      omit the repeated header.
    #   3) Page is later in document order than the best table.
    merged_table = list(best_table['table'])
    pages = [best_table['page']]

    for candidate in candidates[1:]:
        # Skip pages earlier than best_table's page (continuations come after)
        if candidate['page'] < best_table['page']:
            continue
        cand_col_count = len(candidate['header']) if candidate['header'] else 0
        if cand_col_count != best_col_count or cand_col_count == 0:
            continue

        if _headers_match(best_table['header'], candidate['header']):
            # Strip duplicate header, append data rows.
            merged_table.extend(candidate['table'][1:])
            if candidate['page'] not in pages:
                pages.append(candidate['page'])
        elif candidate['score'] >= best_table['score'] * 0.35:
            # Same column count, no header repeat — treat all rows as data.
            # 0.35 multiplier accepts continuation tables that score lower
            # because they have no header keywords on a data-only page.
            merged_table.extend(candidate['table'])
            if candidate['page'] not in pages:
                pages.append(candidate['page'])

    # Convert to DataFrame
    df = pd.DataFrame(merged_table[1:], columns=merged_table[0])

    return {
        'table_data': df,
        'confidence': min(best_table['score'] / 50, 1.0),  # Normalize to 0-1
        'page_numbers': pages,
        'header_row': best_table['header']
    }


def split_merged_cells(df: pd.DataFrame) -> pd.DataFrame:
    """
    Detect and fill vertically merged cells.

    Logic:
    - For each column, scan top-to-bottom
    - If cell is empty AND previous cell had value, propagate value down
    - Stop propagation when non-empty cell found
    - Special case: MPN column should NEVER be merged (indicates multi-line item)

    Args:
        df: DataFrame with potential merged cells

    Returns:
        DataFrame with merged cells split and values propagated
    """
    df = df.copy()

    for col in df.columns:
        # Skip potential MPN columns - empty MPN = multi-line description
        col_lower = str(col).lower()
        if any(keyword in col_lower for keyword in ['mpn', 'part number', 'p/n', 'part #']):
            continue

        last_value = None
        for idx in range(len(df)):
            try:
                current = df.at[idx, col]

                # Handle case where df.at returns a Series (shouldn't happen, but does sometimes)
                if isinstance(current, pd.Series):
                    current = current.iloc[0] if len(current) > 0 else None

                if pd.isna(current) or str(current).strip() == '':
                    if last_value is not None:
                        df.at[idx, col] = last_value  # Propagate
                else:
                    last_value = str(current).strip()
            except (KeyError, IndexError):
                # Skip if column doesn't exist or index out of range
                continue

    return df


def consolidate_multiline_items(df: pd.DataFrame, mpn_col: str, desc_col: str) -> pd.DataFrame:
    """
    Merge runs of rows where MPN is empty (description continuation).

    Improvement over the previous version: handles description spans of any
    length, not just 1 continuation row. Each empty-MPN row that follows a
    valid MPN row appends its description to the parent row's description
    (newline-joined) until we hit another non-empty MPN.
    """
    df = df.copy()
    rows_to_delete: list = []

    def _scalar(v):
        if isinstance(v, pd.Series):
            return v.iloc[0] if len(v) > 0 else None
        return v

    parent_idx = None
    for idx in range(len(df)):
        try:
            current_mpn = _scalar(df.at[idx, mpn_col])
        except (KeyError, IndexError):
            parent_idx = None
            continue

        is_empty = pd.isna(current_mpn) or str(current_mpn).strip() == ''
        if not is_empty:
            parent_idx = idx
            continue

        if parent_idx is None:
            continue

        # This row is a continuation. Merge its description into the parent.
        try:
            current_desc = _scalar(df.at[idx, desc_col])
            current_desc_str = '' if pd.isna(current_desc) else str(current_desc).strip()
            if current_desc_str:
                prev_desc = _scalar(df.at[parent_idx, desc_col])
                prev_desc_str = '' if pd.isna(prev_desc) else str(prev_desc).strip()
                df.at[parent_idx, desc_col] = (
                    f"{prev_desc_str}\n{current_desc_str}" if prev_desc_str else current_desc_str
                )
            rows_to_delete.append(idx)
        except (KeyError, IndexError):
            continue

    if rows_to_delete:
        df = df.drop(rows_to_delete).reset_index(drop=True)

    return df


def remove_duplicate_headers(df: pd.DataFrame, original_header: List) -> pd.DataFrame:
    """
    Remove rows that match the original header pattern.

    Logic:
    - Compare each row to the original header
    - If >= 70% of cells match header keywords (case-insensitive), delete row

    Args:
        df: DataFrame with potential duplicate headers
        original_header: Original header row to match against

    Returns:
        DataFrame with duplicate header rows removed
    """
    header_keywords = set(str(h).lower().strip() for h in original_header if h)
    rows_to_delete = []

    for idx in range(len(df)):
        row_values = set(str(v).lower().strip() for v in df.iloc[idx] if v)

        # Calculate overlap
        overlap = len(header_keywords & row_values)
        match_ratio = overlap / len(header_keywords) if header_keywords else 0

        if match_ratio >= 0.7:  # 70% match = likely header
            rows_to_delete.append(idx)

    if rows_to_delete:
        df = df.drop(rows_to_delete).reset_index(drop=True)

    return df


def align_and_clean(df: pd.DataFrame) -> pd.DataFrame:
    """
    Ensure consistent column count and remove empty rows.

    Logic:
    - Fill NaN/None with empty strings
    - Remove rows where ALL cells are empty
    - Strip whitespace from all cells

    Args:
        df: DataFrame to clean

    Returns:
        Cleaned DataFrame
    """
    # Replace NaN/None with empty strings
    df = df.fillna('')

    # Remove rows where all values are empty
    df = df[~df.apply(lambda row: all(str(v).strip() == '' for v in row), axis=1)]

    # Strip whitespace from all cells
    df = df.map(lambda x: str(x).strip() if pd.notna(x) else '')

    return df.reset_index(drop=True)


def detect_column_type(col_name: str) -> Optional[str]:
    """
    Detect column type from header name.

    Returns one of: 'item_number', 'manufacturer', 'mpn', 'customer_pn',
                    'vendor', 'vendor_pn', 'qty', 'refdes', 'description',
                    'revision', None.

    Backed by ai_helpers.classify_column (fuzzy + priority-substring matching),
    so typos like "Manuf. P/N" or "Mfctr. Part #" still resolve correctly.
    """
    return classify_column(col_name)


def map_to_bom_schema(df: pd.DataFrame, extracted_headers: List) -> pd.DataFrame:
    """
    Map extracted columns to enhanced BOM schema.

    Output Schema (all horizontal columns):
        - Item Number (if available)
        - Manufacturer (if available)
        - Manufacturer Part Number (MPN) - REQUIRED
        - Customer Part Number (if available)
        - Vendor/Supplier (if available)
        - Vendor Part Number (if available)
        - Quantity (default to "1" if not found)
        - Reference Designators (if available)
        - Description (if available)

    Args:
        df: DataFrame with extracted data
        extracted_headers: Original header names from PDF

    Returns:
        DataFrame mapped to enhanced BOM schema

    Raises:
        ValueError: If no MPN column detected
    """
    # Detect column types. When multiple columns map to the same type (common
    # for MPN — "Part Number" + "Mfg Part Number" + "Vendor Part Number"),
    # pick the one whose *actual values* score best against the type's pattern.
    # This is the key fix for PDFs where the header looks right but the column
    # is actually item numbers (1, 2, 3…) or empty.
    candidates_by_type: Dict[str, List[str]] = {}
    for col in df.columns:
        col_type = detect_column_type(col)
        if col_type:
            candidates_by_type.setdefault(col_type, []).append(col)

    column_mapping = {}
    for ctype, cols in candidates_by_type.items():
        if len(cols) == 1:
            column_mapping[ctype] = cols[0]
            continue
        # Multiple candidates → score by content
        best_col = cols[0]
        best_score = -1.0
        for c in cols:
            try:
                score = score_column_content(df[c].tolist(), ctype)
            except Exception:
                score = 0.0
            if score > best_score:
                best_score = score
                best_col = c
        column_mapping[ctype] = best_col

    # Build standardized DataFrame with ALL required columns
    standardized = pd.DataFrame()

    # Item Number (if detected, otherwise generate)
    if 'item_number' in column_mapping:
        standardized['Item Number'] = df[column_mapping['item_number']]
    else:
        standardized['Item Number'] = range(1, len(df) + 1)

    # Manufacturer
    if 'manufacturer' in column_mapping:
        standardized['Manufacturer'] = df[column_mapping['manufacturer']]
    else:
        standardized['Manufacturer'] = ''

    # MPN (REQUIRED)
    if 'mpn' in column_mapping:
        standardized['Manufacturer Part Number (MPN)'] = df[column_mapping['mpn']]
    else:
        # Fallback: try to use the existing excel_tool column detection
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            temp_path = tmp.name
            df.to_excel(temp_path, index=False)

        try:
            _, col_map = find_header_row_and_map(temp_path)
            if col_map.get('mpn'):
                mpn_col_name = col_map['mpn']
                # Ensure the column exists in df - might be integer index
                if mpn_col_name in df.columns:
                    standardized['Manufacturer Part Number (MPN)'] = df[mpn_col_name]
                elif isinstance(mpn_col_name, int) and mpn_col_name < len(df.columns):
                    # Use positional indexing
                    standardized['Manufacturer Part Number (MPN)'] = df.iloc[:, mpn_col_name]
                else:
                    # Try to find by looking at column names
                    mpn_found = False
                    for col in df.columns:
                        if detect_column_type(str(col)) == 'mpn':
                            standardized['Manufacturer Part Number (MPN)'] = df[col]
                            mpn_found = True
                            break
                    if not mpn_found:
                        raise ValueError(ERROR_MESSAGES['no_mpn'])
            else:
                raise ValueError(ERROR_MESSAGES['no_mpn'])
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)

    # Customer Part Number
    if 'customer_pn' in column_mapping:
        standardized['Customer Part Number'] = df[column_mapping['customer_pn']]
    else:
        standardized['Customer Part Number'] = ''

    # Vendor/Supplier
    if 'vendor' in column_mapping:
        standardized['Vendor/Supplier'] = df[column_mapping['vendor']]
    else:
        standardized['Vendor/Supplier'] = ''

    # Vendor Part Number
    if 'vendor_pn' in column_mapping:
        standardized['Vendor Part Number'] = df[column_mapping['vendor_pn']]
    else:
        standardized['Vendor Part Number'] = ''

    # Quantity
    if 'qty' in column_mapping:
        standardized['Quantity'] = df[column_mapping['qty']]
    else:
        standardized['Quantity'] = '1'  # Default quantity

    # Unit of Measure (NEW — preserves UOM where source has one)
    if 'uom' in column_mapping:
        standardized['UOM'] = df[column_mapping['uom']]

    # Reference Designators
    if 'refdes' in column_mapping:
        standardized['Reference Designators'] = df[column_mapping['refdes']]
    else:
        standardized['Reference Designators'] = ''

    # Description
    if 'description' in column_mapping:
        standardized['Description'] = df[column_mapping['description']]
    else:
        standardized['Description'] = ''

    # Flag rows with alternate-MPN cells so the comparison tool can split
    # them into primary + alternates downstream.
    try:
        from ai_helpers import parse_alternate_mpns
        alt_flags = standardized['Manufacturer Part Number (MPN)'].apply(
            lambda v: '|'.join(parse_alternate_mpns(v)[1:]) if len(parse_alternate_mpns(v)) > 1 else ''
        )
        if alt_flags.astype(bool).any():
            standardized['Alternate MPNs'] = alt_flags
    except Exception:
        pass

    # Clean up: remove rows with empty MPN
    standardized = standardized[standardized['Manufacturer Part Number (MPN)'].astype(str).str.strip() != '']

    return standardized


def generate_clean_excel(df: pd.DataFrame, output_path: str) -> str:
    """
    Generate Excel file with strict formatting rules.

    Requirements:
    - No merged cells
    - Consistent headers in row 1
    - Data starts at row 2
    - Auto-adjusted column widths
    - Plain text format (no formulas)

    Args:
        df: DataFrame to write to Excel
        output_path: Path where Excel file should be saved

    Returns:
        Path to generated Excel file
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "BOM"

    # Write headers (row 1)
    for col_idx, col_name in enumerate(df.columns, start=1):
        cell = ws.cell(row=1, column=col_idx)
        cell.value = col_name
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=False)

    # Write data (starting row 2)
    for row_idx, row_data in enumerate(df.values, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = str(value) if pd.notna(value) else ''
            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter

        for cell in column:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass

        adjusted_width = min(max_length + 2, 50)  # Cap at 50 chars
        ws.column_dimensions[column_letter].width = adjusted_width

    # Verify no merged cells (should be empty set)
    assert len(ws.merged_cells.ranges) == 0, "Generated Excel contains merged cells!"

    # Save
    wb.save(output_path)
    return output_path


def _write_sheet(ws, df: pd.DataFrame) -> None:
    """Write a DataFrame into an existing worksheet: bold header row 1, plain
    text data from row 2, wrapped cells, auto-width, zero merged cells."""
    for col_idx, col_name in enumerate(df.columns, start=1):
        cell = ws.cell(row=1, column=col_idx)
        cell.value = str(col_name)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=False)

    for row_idx, row_data in enumerate(df.values, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = str(value) if pd.notna(value) else ''
            cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)

    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except Exception:
                pass
        ws.column_dimensions[column_letter].width = min(max_length + 2, 50)


def generate_two_tab_workbook(raw_df: pd.DataFrame, bom_df: pd.DataFrame, output_path: str,
                              images: Optional[List[bytes]] = None) -> str:
    """
    Write the converter output as a two-sheet workbook (fallback when the
    stakeholder template can't be used): "Normal" (raw) + "BOM" (normalized),
    BOM selected. Adds a "Drawing" sheet when page images are provided.
    """
    wb = Workbook()

    normal_ws = wb.active
    normal_ws.title = "Normal"
    _write_sheet(normal_ws, raw_df)

    bom_ws = wb.create_sheet("BOM")
    _write_sheet(bom_ws, bom_df if bom_df is not None and not bom_df.empty else raw_df)

    _add_drawing_sheet(wb, images)

    # Make BOM the selected/active tab on open.
    wb.active = wb.index(bom_ws)
    for sheet in wb.worksheets:
        sheet.sheet_view.tabSelected = (sheet is bom_ws)

    wb.save(output_path)
    return output_path


# Stakeholder quote template bundled with the backend; BOM data is written into
# its line-item area so the output matches the format buyers already use.
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "BOM_QUOTE_TEMPLATE.xlsx")
_TEMPLATE_HEADER_ROW = 23  # row 24 is the first data row; columns A..G = QUOTE_TEMPLATE_COLUMNS


def generate_template_bom_workbook(bom_df: pd.DataFrame, raw_df: pd.DataFrame, output_path: str,
                                   images: Optional[List[bytes]] = None) -> str:
    """
    Fill a copy of BOM QUOTE TEMPLATE.xlsx with the extracted BOM rows so the
    output matches the stakeholder's quoting format (title block, headers,
    pricing-tier columns and formulas all preserved). The BOM line items are
    written into columns A–G starting at row 24. A "Normal (raw)" sheet and a
    "Drawing" sheet are appended. Raises if the template can't be loaded so the
    caller can fall back to the plain two-tab workbook.
    """
    from openpyxl import load_workbook as _load_wb

    wb = _load_wb(TEMPLATE_PATH)  # keep formulas/styling
    ws = wb["BOM"] if "BOM" in wb.sheetnames else wb.active

    cols = list(bom_df.columns) if bom_df is not None else []
    # Map our bom_df columns onto template columns A..G (QUOTE_TEMPLATE_COLUMNS).
    def _col_for(label):
        return cols.index(label) if label in cols else None
    idx = {c: _col_for(c) for c in QUOTE_TEMPLATE_COLUMNS}

    start = _TEMPLATE_HEADER_ROW + 1
    n = 0 if bom_df is None else len(bom_df)
    for i in range(n):
        row = bom_df.iloc[i]
        for c_off, label in enumerate(QUOTE_TEMPLATE_COLUMNS):  # A..G
            j = idx.get(label)
            val = "" if j is None else row.iloc[j]
            ws.cell(row=start + i, column=1 + c_off).value = ("" if pd.isna(val) else str(val))

    # Append the raw extraction + drawing as extra sheets.
    raw_ws = wb.create_sheet("Normal (raw)")
    _write_sheet(raw_ws, raw_df)
    _add_drawing_sheet(wb, images)

    # Open on the BOM sheet.
    try:
        wb.active = wb.index(ws)
        for sheet in wb.worksheets:
            sheet.sheet_view.tabSelected = (sheet is ws)
    except Exception:
        pass

    wb.save(output_path)
    return output_path


def render_drawing_images(pdf_path: str, max_pages: int = 3, max_px: int = 1600) -> List[bytes]:
    """Render the first pages of the PDF to downscaled PNGs for embedding as a
    visual reference in the output. Best-effort: returns [] if pdf2image/Pillow
    or poppler are unavailable, or rendering fails — never raises."""
    try:
        from pdf2image import convert_from_path
        from PIL import Image as PILImage  # noqa: F401
    except Exception:
        return []
    try:
        pages = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=max_pages)
    except Exception:
        return []
    out: List[bytes] = []
    for pg in pages:
        try:
            w, h = pg.size
            scale = min(1.0, float(max_px) / max(w, h))
            if scale < 1.0:
                pg = pg.resize((max(1, int(w * scale)), max(1, int(h * scale))))
            buf = io.BytesIO()
            pg.save(buf, format="PNG")
            out.append(buf.getvalue())
        except Exception:
            continue
    return out


def _add_drawing_sheet(wb, images: Optional[List[bytes]]):
    """Add a 'Drawing' sheet with the rendered page image(s) stacked vertically."""
    if not images:
        return
    try:
        from openpyxl.drawing.image import Image as XLImage
    except Exception:
        return
    ws = wb.create_sheet("Drawing")
    ws["A1"] = "Source drawing (visual reference)"
    ws["A1"].font = Font(bold=True)
    row = 3
    keep = []  # hold BytesIO refs alive until wb.save reads them
    for png in images:
        try:
            bio = io.BytesIO(png)
            keep.append(bio)
            img = XLImage(bio)
            ws.add_image(img, f"A{row}")
            row += int((getattr(img, "height", 600) or 600) / 18) + 4
        except Exception:
            continue
    ws._bom_keep_images = keep  # prevent GC before save


def generate_single_sheet_workbook(df: pd.DataFrame, output_path: str, sheet_name: str = "PDF",
                                   images: Optional[List[bytes]] = None) -> str:
    """Write a one-sheet workbook (used by 'normal' mode — any PDF → raw table)."""
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "PDF"
    _write_sheet(ws, df)
    assert len(ws.merged_cells.ranges) == 0, "Generated Excel contains merged cells!"
    _add_drawing_sheet(wb, images)
    wb.save(output_path)
    return output_path


def generate_word_document(frames: List, output_path: str,
                           images: Optional[List[bytes]] = None) -> str:
    """
    Write a .docx containing one heading + table per (heading, DataFrame) pair,
    then (if provided) a "Drawing" section with the rendered page image(s).

    Used for the Word output option. Empty frames are skipped. Cells are plain
    text (wrapped text already merged upstream), so there are no merged cells.
    """
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    wrote_any = False
    for heading, df in frames:
        if df is None or len(df) == 0:
            continue
        doc.add_heading(str(heading), level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        for i, col in enumerate(df.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = "" if pd.isna(val) else str(val)
        doc.add_paragraph("")
        wrote_any = True

    if not wrote_any:
        doc.add_paragraph("No table rows were extracted from this PDF.")

    if images:
        doc.add_heading("Drawing", level=1)
        for png in images:
            try:
                doc.add_picture(io.BytesIO(png), width=Inches(6.5))
            except Exception:
                continue

    doc.save(output_path)
    return output_path


def _docx_shade(cell, fill_hex: str):
    """Set a cell's background fill (mimics the template header shading)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def generate_word_bom_document(bom_df: pd.DataFrame, raw_df: pd.DataFrame, output_path: str,
                               images: Optional[List[bytes]] = None) -> str:
    """
    Word output styled to match BOM QUOTE TEMPLATE.xlsx: a "BOM QUOTE" title, a
    summary block (Total / SMT / Through-Hole / Fine-Pitch counts computed from
    the data), and the 7-column line-item table with the template's column widths
    and a shaded, bold header. Adds the raw extraction + drawing image after.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cols = QUOTE_TEMPLATE_COLUMNS  # Item# | Description | Mfg | Mfg P/N | Qty | SMT/TH | Loc
    widths = [0.5, 3.4, 1.5, 1.6, 0.7, 0.85, 1.0]  # inches, proportional to the template

    def _int(v):
        try:
            return int(float(str(v).strip()))
        except Exception:
            return 0

    qty = list(bom_df["Qty"]) if "Qty" in bom_df.columns else []
    smt = list(bom_df["SMT/TH"]) if "SMT/TH" in bom_df.columns else [""] * len(bom_df)
    desc = list(bom_df["Description"]) if "Description" in bom_df.columns else [""] * len(bom_df)
    total_comp = sum(_int(q) for q in qty)
    smt_comp = sum(_int(q) for q, s in zip(qty, smt) if str(s).upper() in ("SMT", "SMD"))
    th_comp = sum(_int(q) for q, s in zip(qty, smt) if str(s).upper() == "TH")
    fp_comp = sum(_int(q) for q, d in zip(qty, desc)
                  if any(k in str(d).upper() for k in ("QFN", "BGA", "QFP", "FINE PITCH")))

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BOM QUOTE")
    run.bold = True
    run.font.size = Pt(18)

    # Summary block
    summ = doc.add_table(rows=0, cols=2)
    for label, val in (("Total Components", total_comp), ("Lines", len(bom_df)),
                       ("SMT", smt_comp), ("Through Hole", th_comp), ("Fine Pitch/QFP/QFN/BGA", fp_comp)):
        r = summ.add_row().cells
        r[0].text = label
        r[1].text = str(val)
        for p in r[0].paragraphs:
            for rn in p.runs:
                rn.bold = True
    doc.add_paragraph("")

    # Line-item table
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
        _docx_shade(hdr[i], "D9E1F2")  # light blue, like a quote header
        for p in hdr[i].paragraphs:
            for rn in p.runs:
                rn.bold = True
    for _, row in bom_df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            v = row[c] if c in bom_df.columns else ""
            cells[i].text = "" if pd.isna(v) else str(v)
    # Apply column widths to every cell (python-docx needs per-cell width).
    for r in table.rows:
        for i, w in enumerate(widths):
            r.cells[i].width = Inches(w)

    # Raw extraction
    doc.add_paragraph("")
    doc.add_heading("Normal (raw extraction)", level=2)
    if raw_df is not None and len(raw_df) > 0:
        rt = doc.add_table(rows=1, cols=len(raw_df.columns))
        rt.style = "Table Grid"
        for i, c in enumerate(raw_df.columns):
            rt.rows[0].cells[i].text = str(c)
            for p in rt.rows[0].cells[i].paragraphs:
                for rn in p.runs:
                    rn.bold = True
        for _, row in raw_df.iterrows():
            cells = rt.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = "" if pd.isna(val) else str(val)

    if images:
        doc.add_heading("Drawing", level=2)
        for png in images:
            try:
                doc.add_picture(io.BytesIO(png), width=Inches(9.5))
            except Exception:
                continue

    doc.save(output_path)
    return output_path


# Column layout used on the "BOM" tab — mirrors the line-item header (row 23)
# of BOM QUOTE TEMPLATE.xlsx so the result pastes straight into the stakeholder
# quoting workbook. "Loc" carries reference designators; "SMT/TH" (assembly
# type) is left blank because it cannot be inferred from a BOM PDF.
QUOTE_TEMPLATE_COLUMNS = ["Item#", "Description", "Mfg", "Mfg P/N", "Qty", "SMT/TH", "Loc"]

# Package/footprint hints for inferring assembly type (SMT vs through-hole).
_TH_HINTS = ("THRU", "THROUGH", "PTH", "DIP", "PDIP", "RADIAL", "AXIAL",
             "TO-220", "TO220", "TO-92", "TO92", "TO-247", "TO-218")
_SMT_HINTS = ("SURFACE MNT", "SURFACE MOUNT", "SMD", "SMT", "0201", "0402", "0603", "0805", "1206", "1210",
              "1806", "2010", "2512", "SOT23", "SOT-23", "SOT", "SOIC", "SO-8",
              "SO8", "TSSOP", "SSOP", "MSOP", "TSOP", "QFN", "QFP", "TQFP",
              "LQFP", "BGA", "WSON", "DFN", "SC70", "SC-70", "LGA", "PLCC",
              "MELF", "TUMD", "GS08", "PPAK", "WLCSP", "DPAK", "D2PAK",
              "TO-252", "TO-263", "CHIPLED", "WSON", "0603")


# Non-placed items (the bare board and process materials) have no SMT/TH.
_NON_PLACED = ("PWB", "PCB", "BARE BOARD", "EPOXY", "ADHESIVE", "SOLDER",
               "CONFORMAL", "COATING", "COMPOUND", "INK", "LABEL", "STENCIL", "STAKING")


def _infer_smt_th(text: str) -> str:
    """Infer 'SMT' or 'TH' for a part. Bare board / materials → '' (N/A);
    explicit through-hole packages → 'TH'; otherwise default to 'SMT' (placed
    components and connectors on an SMT assembly), which the quoter can adjust."""
    u = (text or "").upper()
    if any(k in u for k in _NON_PLACED):
        return ""
    if any(h in u for h in _TH_HINTS):
        return "TH"
    if any(h in u for h in _SMT_HINTS):
        return "SMT"
    return "SMT"


# Engineering-drawing header phrases the general classifier doesn't know.
# MIL-STD-100 / ASME Y14.34 parts lists label the part-number column
# "PART OR IDENTIFYING NO." and the CAGE column "CODE IDENT". Checked as a
# fallback only — local to the converter so the shared comparison classifier
# (ai_helpers) is untouched.
_ENGINEERING_HEADER_HINTS = [
    ("part or identifying", "mpn"),
    ("part identifying", "mpn"),
    ("identifying no", "mpn"),
    ("identifying number", "mpn"),
    ("nomenclature", "description"),
    ("qty reqd", "qty"),
    ("qty req", "qty"),
    ("reference designation", "refdes"),
]


def _detect_col_type_eng(col) -> Optional[str]:
    """detect_column_type() with an engineering-drawing header fallback."""
    ctype = detect_column_type(str(col))
    if ctype:
        return ctype
    norm = str(col).strip().lower()
    for needle, t in _ENGINEERING_HEADER_HINTS:
        if needle in norm:
            return t
    return None


def _pick_best_columns(df: pd.DataFrame) -> Dict[str, str]:
    """Map each canonical type to its best source column. When several columns
    share a type (e.g. two "Part Number" columns), pick the one whose *values*
    score highest for that type — header is a hint, content is proof."""
    candidates: Dict[str, List[str]] = {}
    for col in df.columns:
        ctype = _detect_col_type_eng(col)
        if ctype:
            candidates.setdefault(ctype, []).append(col)

    chosen: Dict[str, str] = {}
    for ctype, cols in candidates.items():
        if len(cols) == 1:
            chosen[ctype] = cols[0]
            continue
        best_col, best_score = cols[0], -1.0
        for c in cols:
            try:
                score = score_column_content(df[c].tolist(), ctype)
            except Exception:
                score = 0.0
            if score > best_score:
                best_score, best_col = score, c
        chosen[ctype] = best_col
    return chosen


def map_to_quote_template_schema(df: pd.DataFrame) -> pd.DataFrame:
    """
    Map an extracted table to the BOM QUOTE TEMPLATE column layout
    (Item# | Description | Mfg | Mfg P/N | Qty | SMT/TH | Loc).

    This is the "BOM" tab. The lossless raw table lives on the "Normal" tab, so
    here we keep only the stakeholder columns and drop rows that carry neither a
    part number nor a description (header echoes, note lines, blanks).
    """
    chosen = _pick_best_columns(df)

    out = pd.DataFrame()
    out["Item#"] = (
        df[chosen["item_number"]].astype(str)
        if "item_number" in chosen else [str(i) for i in range(1, len(df) + 1)]
    )
    out["Description"] = df[chosen["description"]].astype(str) if "description" in chosen else ""
    out["Mfg"] = df[chosen["manufacturer"]].astype(str) if "manufacturer" in chosen else ""
    out["Mfg P/N"] = df[chosen["mpn"]].astype(str) if "mpn" in chosen else ""
    out["Qty"] = df[chosen["qty"]].astype(str) if "qty" in chosen else ""
    # SMT/TH isn't an explicit column on the drawing — infer it from each part's
    # package/footprint in the description + part number.
    mpn_series = df[chosen["mpn"]].astype(str) if "mpn" in chosen else pd.Series([""] * len(out))
    out["SMT/TH"] = [
        _infer_smt_th(f"{d} {p}") for d, p in zip(out["Description"], list(mpn_series))
    ]
    out["Loc"] = df[chosen["refdes"]].astype(str) if "refdes" in chosen else ""

    # Keep only real line items: must have a part number or a description.
    def _has(v) -> bool:
        s = str(v).strip().lower()
        return s not in ("", "nan", "none")

    mask = out["Mfg P/N"].apply(_has) | out["Description"].apply(_has)
    out = out[mask].reset_index(drop=True)
    # Renumber Item# sequentially if the source had no usable item column.
    if "item_number" not in chosen:
        out["Item#"] = [str(i) for i in range(1, len(out) + 1)]
    return out


def _extract_generic_tables_from_pdf(pdf_path: str) -> Optional[Dict]:
    """Generic fallback: pull whatever table(s) pdfplumber finds, no BOM scoring."""
    all_tables: List[List[List[str]]] = []
    pages: List[int] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            for table in page.extract_tables() or []:
                if table and len(table) >= 2 and _is_table_data_valid(table):
                    all_tables.append(table)
                    pages.append(page_num + 1)
    if not all_tables:
        return None
    primary = all_tables[0]
    df = pd.DataFrame(primary[1:], columns=[
        (str(c).strip() if c is not None else f"Column {i + 1}")
        for i, c in enumerate(primary[0])
    ])
    # De-duplicate blank/repeated column names.
    seen: Dict[str, int] = {}
    new_cols: List[str] = []
    for i, c in enumerate(df.columns):
        n = c or f"Column {i + 1}"
        seen[n] = seen.get(n, 0) + 1
        if seen[n] > 1:
            n = f"{n} ({seen[n]})"
        new_cols.append(n)
    df.columns = new_cols
    return {
        'table_data': df,
        'header_row': list(df.columns),
        'confidence': 0.5,  # unknown — mark as medium
        'page_numbers': pages[:1],
    }


def extract_bom_from_pdf(pdf_path: str, output_excel_path: str,
                         mode: str = "bom", out_format: str = "excel",
                         progress_cb=None, text_model: Optional[str] = None) -> Dict:
    """
    Main entry point: extract a table from a PDF and write the output file.

    mode:
        "bom"    — BOM-aware: normalize to the BOM/quote schema, AI vision
                   fallback for dense drawings. Excel output is two sheets
                   ("Normal" raw + "BOM", BOM first); Word output is both tables.
        "normal" — generic any-PDF table extraction, raw, single sheet/table.
    out_format:
        "excel" (.xlsx) or "word" (.docx).

    Returns a dict whose 'preview' holds the rows/columns the UI should show
    (the BOM table for bom mode, the raw table for normal mode).
    """
    warnings = []
    is_bom_like = True
    mode = (mode or "bom").lower()
    out_format = (out_format or "excel").lower()

    def _p(pct, stage):
        if progress_cb:
            try:
                progress_cb(pct, stage)
            except Exception:
                pass

    # Step 1: Validate digital PDF
    _p(5, "Reading PDF")
    table_data = None
    early_layout = False
    is_scanned = False
    try:
        validate_digital_pdf(pdf_path)
    except ValueError as e:
        # Image-only/scanned PDF: don't give up — OCR is the only way in. Other
        # validation failures (encrypted, corrupt) still propagate.
        if "scanned" in str(e).lower() or str(e) == ERROR_MESSAGES['scanned_pdf']:
            is_scanned = True
        else:
            raise

    if is_scanned and mode == "bom":
        # Best-effort OCR extraction (cable/ground-strap & callout drawings).
        # Return what OCR finds with a loud verify flag, or a clear message.
        _p(20, "OCR (scanned PDF)")
        harn_df = None
        try:
            from harness_parse import parse_harness_bom
            harn_df = parse_harness_bom(pdf_path)
        except Exception:
            harn_df = None
        if harn_df is None or len(harn_df) < 2:
            raise ValueError(
                "This PDF is scanned/image-only and OCR could not extract a BOM. "
                "Please provide a digitally generated PDF with selectable text."
            )
        table_data = {
            'table_data': harn_df,
            'header_row': list(harn_df.columns),
            'confidence': 0.4,
            'page_numbers': [1],
            'extraction_method': 'ocr-scanned',
        }
        early_layout = True
        warnings.append(
            "Scanned/image-only PDF — extracted via OCR. Verify ALL part numbers "
            "and descriptions against the drawing."
        )

    # Step 2: Detect the BOM table.

    # Fast path (bom mode): try the deterministic columnar parser first. On
    # engineering drawings this returns the parts list in <1s and lets us skip
    # the slow geometry scan. It returns None for non-columnar BOMs, which then
    # fall through to the normal detection below.
    if mode == "bom" and table_data is None:
        try:
            from layout_parse import parse_bom
            _p(12, "Parsing parts-list columns")
            cand = parse_bom(pdf_path)
            if cand is not None and len(cand) >= 5:
                table_data = {
                    'table_data': cand,
                    'header_row': list(cand.columns),
                    'confidence': 0.9,
                    'page_numbers': [1],
                    'extraction_method': 'layout-columns',
                }
                early_layout = True
                warnings.append("Parsed the drawing's parts-list columns directly.")
        except Exception:
            table_data = None

    if table_data is None:
        _p(15, "Detecting tables")
        try:
            table_data = detect_bom_table(pdf_path)
        except ValueError as e:
            # Grid detection failed. Try the layout parsers (item-anchored
            # drawing tables + wide gridless PCB/"Part Information" exports)
            # before giving up — these read clean text-layer columns that have
            # no gridlines for pdfplumber to find.
            lay_df = None
            try:
                from layout_parse import parse_bom as _layout_parse_bom
                _p(18, "Parsing text-layer columns")
                lay_df = _layout_parse_bom(pdf_path, include_generic=True)
            except Exception:
                lay_df = None
            if lay_df is not None and len(lay_df) >= 3:
                table_data = {
                    'table_data': lay_df,
                    'header_row': list(lay_df.columns),
                    'confidence': 0.85,
                    'page_numbers': [1],
                    'extraction_method': 'layout-columns',
                }
                early_layout = True
            else:
                fallback = _extract_generic_tables_from_pdf(pdf_path)
                if fallback is None:
                    raise ValueError(
                        "No table detected in this PDF. Ensure the PDF contains a table with "
                        "a header row and at least one data row."
                    ) from e
                warnings.append(f"BOM detection skipped: {e}. Exporting raw table.")
                table_data = fallback
                is_bom_like = False

    if table_data['confidence'] < 0.6:
        warnings.append(f"Table detected with low confidence ({table_data['confidence']:.0%}). Results may require manual review.")

    # Keep the source's own columns and order — no BOM schema expansion.
    df = table_data['table_data']

    if is_bom_like and not early_layout:
        df = split_merged_cells(df)
        df = remove_duplicate_headers(df, table_data['header_row'])
        df = align_and_clean(df)
        # Multi-line description merging (>1 continuation row supported now).
        try:
            mpn_col = next((c for c in df.columns if detect_column_type(c) == 'mpn'), None)
            desc_col = next((c for c in df.columns if detect_column_type(c) == 'description'), None)
            if mpn_col and desc_col:
                df = consolidate_multiline_items(df, mpn_col, desc_col)
        except Exception:
            pass

    extraction_method = table_data.get('extraction_method', 'tables')

    raw_df = df
    bom_df = None

    if mode == "bom":
        # BOM-aware output. First map the text-extracted table to the BOM schema.
        def _safe_map(d):
            try:
                return map_to_quote_template_schema(d)
            except Exception:
                return None

        _p(20, "Mapping BOM columns")
        bom_df = _safe_map(df)

        # The geometry parser often *looks* like it found a table on a dense
        # engineering drawing (it grabs the title block) yet yields no real BOM
        # rows. The reliable failure signal is "few/zero mapped BOM rows".
        text_bom_rows = 0 if bom_df is None else len(bom_df)
        if text_bom_rows < 3 or table_data['confidence'] < 0.6:
            # 0) Fast deterministic parse of columnar drawing parts lists
            #    (pdftotext -layout keeps the columns aligned; the CAGE code
            #    anchors each row). Instant and free — handles MIL-STD/ASME
            #    "LIST OF MATERIALS" drawings the geometry parser garbles.
            try:
                from layout_parse import parse_bom
                _p(25, "Parsing parts-list columns")
                # Last-resort layout parse: the grid detector already failed
                # here, so allow the broader generic item-anchored parser too.
                lay_df = parse_bom(pdf_path, include_generic=True)
                if lay_df is not None and not lay_df.empty:
                    lay_bom = _safe_map(lay_df)
                    if lay_bom is not None and len(lay_bom) > text_bom_rows:
                        raw_df = lay_df
                        bom_df = lay_bom
                        text_bom_rows = len(lay_bom)
                        extraction_method = 'layout-columns'
                        warnings.append("Parsed the drawing's parts-list columns directly.")
            except Exception as e:
                warnings.append(f"Layout parse error: {e}")

        # 0.5) Cable / wire-harness drawings: the BOM is callouts, not a table.
        #      Pull connector/pin/wire callouts directly (text layer, or OCR when
        #      the part numbers were flattened to vector outlines). Deterministic
        #      and fast — runs before the slow LLM/vision fallbacks.
        if text_bom_rows < 3:
            try:
                from harness_parse import parse_harness_bom
                _p(28, "Reading harness callouts")
                harn_df = parse_harness_bom(pdf_path)
                if harn_df is not None and not harn_df.empty:
                    harn_bom = _safe_map(harn_df)
                    if harn_bom is not None and len(harn_bom) > text_bom_rows:
                        raw_df = harn_df
                        bom_df = harn_bom
                        text_bom_rows = len(harn_bom)
                        extraction_method = 'harness-callouts'
                        warnings.append("Extracted harness/cable callouts from the drawing.")
            except Exception as e:
                warnings.append(f"Harness parse error: {e}")

        if (text_bom_rows < 3 or table_data['confidence'] < 0.6) and \
                os.environ.get("BOM_TEXT_LLM_ENABLED", "false").lower() in ("1", "true", "yes"):
            # 1) Slower opt-in fallback: hand the PDF text layer to a local text
            #    LLM. Off by default — slow on CPU and the deterministic layout
            #    parser already handles the common columnar drawings.
            try:
                from text_extract_llm import extract_bom_via_text_llm, text_llm_available
                if text_llm_available():
                    _p(30, "AI structuring BOM from text")
                    llm_df = extract_bom_via_text_llm(pdf_path, progress_cb=progress_cb, model=text_model)
                    if llm_df is not None and not llm_df.empty:
                        llm_bom = _safe_map(llm_df)
                        if llm_bom is not None and len(llm_bom) > text_bom_rows:
                            raw_df = llm_df
                            bom_df = llm_bom
                            text_bom_rows = len(llm_bom)
                            extraction_method = 'ollama-text-llm'
                            warnings.append("Used local AI text model to structure the parts list.")
            except Exception as e:
                warnings.append(f"AI text-structuring error: {e}")

            # 2) Optional slow fallback: image vision. Off by default because on
            #    a CPU host it's slow and tiling can miss the parts-list region.
            #    Enable with BOM_VISION_ENABLED=true.
            if text_bom_rows < 3 and os.environ.get("BOM_VISION_ENABLED", "false").lower() in ("1", "true", "yes"):
                try:
                    from vision_tool import extract_bom_via_vision, vision_available
                    if vision_available():
                        _p(40, "AI vision: rendering drawing")
                        vision_df = extract_bom_via_vision(pdf_path, progress_cb=progress_cb)
                        if vision_df is not None and not vision_df.empty:
                            vision_bom = _safe_map(vision_df)
                            if vision_bom is not None and len(vision_bom) > text_bom_rows:
                                raw_df = vision_df
                                bom_df = vision_bom
                                extraction_method = 'ollama-vision'
                                warnings.append("Used local AI vision (Ollama).")
                except Exception as e:
                    warnings.append(f"AI vision fallback error: {e}")

        if bom_df is None:
            bom_df = raw_df
            warnings.append("BOM normalization produced no rows; BOM tab mirrors the raw table.")

        # When a dedicated extractor (drawing parts-list, harness callouts, AI)
        # actually produced a BOM, the earlier grid-detection notes are noise —
        # they describe a step we recovered from. Replace them with one clear
        # note and a confidence that reflects the method that succeeded, so the
        # UI doesn't flash a scary "detection failed / 50%" on a good result.
        _ALT_CONF = {
            'harness-callouts': 0.85, 'layout-columns': 0.9,
            'ollama-text-llm': 0.8, 'ollama-vision': 0.8,
        }
        if extraction_method in _ALT_CONF and bom_df is not None and len(bom_df) >= 3:
            table_data['confidence'] = max(table_data.get('confidence', 0), _ALT_CONF[extraction_method])
            _noise = (
                'could not detect a bom table', 'low confidence', 'exporting raw table',
                'bom detection skipped', 'bom normalization produced no rows',
                'extracted harness', "parsed the drawing's parts-list columns",
            )
            warnings[:] = [w for w in warnings if not any(n in w.lower() for n in _noise)]
            _label = {
                'harness-callouts': "Extracted the cable/harness callouts (connectors, pins, wires, materials) from the drawing — verify part numbers against the drawing.",
                'layout-columns': "Parsed the drawing's parts-list table directly.",
                'ollama-text-llm': "Structured the parts list with the local AI text model.",
                'ollama-vision': "Read the drawing with local AI vision.",
            }[extraction_method]
            if _label not in warnings:
                warnings.insert(0, _label)

        preview_df = bom_df if (bom_df is not None and not bom_df.empty) else raw_df
        _p(90, "Rendering drawing image")
        images = render_drawing_images(pdf_path)
        _p(94, "Writing file")
        if out_format == "word":
            generate_word_bom_document(bom_df, raw_df, output_excel_path, images=images)
            sheets = ["BOM (quote)", "Normal (raw)", "Drawing"]
        else:
            # Fill the stakeholder quote template; fall back to the plain
            # two-tab workbook if the template can't be loaded.
            try:
                generate_template_bom_workbook(bom_df, raw_df, output_excel_path, images=images)
                sheets = ["BOM (quote template)", "Normal (raw)", "Drawing"]
            except Exception as e:
                warnings.append(f"Quote-template export failed, used plain layout: {e}")
                generate_two_tab_workbook(raw_df, bom_df, output_excel_path, images=images)
                sheets = ["Normal", "BOM", "Drawing"]
    else:
        # Normal: any PDF → raw table, single sheet/table, no BOM normalization.
        preview_df = raw_df
        _p(55, "Rendering drawing image")
        images = render_drawing_images(pdf_path)
        _p(60, "Writing file")
        if out_format == "word":
            generate_word_document([("PDF", raw_df)], output_excel_path, images=images)
            sheets = ["PDF", "Drawing"]
        else:
            generate_single_sheet_workbook(raw_df, output_excel_path, sheet_name="PDF", images=images)
            sheets = ["PDF", "Drawing"]
    _p(99, "Finalizing")

    # Per-column confidence — content-pattern score per detected type. Lets
    # the UI flag specific weak columns ("Qty: 60% — review qty mapping")
    # instead of a single global confidence.
    column_confidence: Dict[str, float] = {}
    try:
        from ai_helpers import score_column_content
        for col in df.columns:
            ctype = detect_column_type(str(col))
            if ctype:
                try:
                    column_confidence[str(col)] = round(score_column_content(df[col].tolist(), ctype), 2)
                except Exception:
                    pass
    except Exception:
        pass

    # Preview the table the user actually asked for: the BOM table in bom mode,
    # the raw table in normal mode. Works for both Excel and Word output.
    preview_cols = [str(c) for c in preview_df.columns]
    preview_rows = [["" if pd.isna(v) else str(v) for v in row] for row in preview_df.values]

    return {
        'success': True,
        'excel_path': output_excel_path,
        'preview': {
            'columns': preview_cols,
            'rows': preview_rows,
            'total_rows': len(preview_df),
        },
        'metadata': {
            'pages_processed': table_data['page_numbers'],
            'rows_extracted': len(df),
            'confidence': table_data['confidence'],
            'columns_detected': list(df.columns),
            'column_confidence': column_confidence,
            'extraction_method': extraction_method,
            'mode': mode,
            'format': out_format,
            'bom_rows': len(bom_df) if bom_df is not None else 0,
            'sheets': sheets,
        },
        'warnings': warnings
    }
